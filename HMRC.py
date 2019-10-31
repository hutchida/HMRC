
import sys
import pandas as pd
import numpy as np
import time
import timeit
import httplib2
import requests
import xml.etree.ElementTree as ET
from lxml import etree
from bs4 import BeautifulSoup, SoupStrainer
import csv
import os
import glob
import datetime
import re
import string
import difflib
from difflib import ndiff
from docx import Document
import win32com.client

if sys.version_info[0] < 3:
    from StringIO import StringIO
else:
    from io import StringIO

pd.set_option('display.max_colwidth', -1) #stop the dataframe from truncating cell contents.

def DirectoryComparison(logpath, xmldir1, xmldir2):
    print('Building xml list1...')
    df1 = pd.DataFrame({'col':os.listdir(xmldir1)})
    df1.to_csv(logpath+'old.csv', sep=',',index=False)
    print('Building xml list2...')
    df2 = pd.DataFrame({'col':os.listdir(xmldir2)})
    df2.to_csv(logpath+'new.csv', index=False)
    
    
    df1 = pd.read_csv(logpath+'old.csv')
    df2 = pd.read_csv(logpath+'new.csv')

    dfdel = pd.DataFrame(columns=['col'])
    for index, item in df1.itertuples():  #loop through each item in the older list
        if any(df2.col == item) == False: #if the older list item doesn't appear in the newer list, it must have been deleted
            dfdel.loc[len(dfdel)] = [item] #add item to end of fresh dataframe
            
    dfdel.to_csv(logpath+'deletions.csv', index=False)
    print('Exported deletions...')
    
    dfadd = pd.DataFrame(columns=['col'])
    for index, item in df2.itertuples(): #loop through each item in the newer list
        if any(df1.col == item) == False: #if the newer list item doesn't appear in the older list, it must have been added
            dfadd.loc[len(dfadd)] = [item] #add item to end of fresh dataframe
    
    dfadd.to_csv(logpath+'additions.csv', index=False)
    print('Exported additions...')
    
    dfrem1 = pd.DataFrame(columns=['col'])
    for index, item in df1.itertuples(): #loop through each item in the older list
        if any(dfdel.col == item) == False: #if the older item doesn't appear in the deleted list, must appear on both new and old lists
            dfrem1.loc[len(dfrem1)] = [item] #add item to end of fresh dataframe
    
    dfrem1.to_csv(logpath+'rem1.csv', index=False)
    print('Exported remainder version 1...')
    
    dfrem2 = pd.DataFrame(columns=['col'])
    for index, item in df2.itertuples(): #loop through each item in the older list
        if any(dfadd.col == item) == False: #if the older item doesn't appear in the deleted list, must appear on both new and old lists
            dfrem2.loc[len(dfrem2)] = [item] #add item to end of fresh dataframe
    
    dfrem2.to_csv(logpath+'rem2.csv', index=False) #rem1 and rem2 should be identical
    print('Exported remainder version 2...')
    
    print("\nTime taken so far..." + str(datetime.datetime.now() - start))


def Prelim(xmldir1, xmldir2, logpath, start):
    
    print('Looping through list of potential changed docs to compare byte size...')
    dfrem = pd.read_csv(logpath+'rem1.csv')
    dfchange = pd.DataFrame(columns=['col'])
    i = 0
    for index, item in dfrem.itertuples(): #loop through each item in the remainder list
        while True:
            try:
                oldbytecount = os.path.getsize(xmldir1 + item)
                newbytecount = os.path.getsize(xmldir2 + item)
                if oldbytecount != newbytecount: 
                    dfchange.loc[len(dfchange)] = [item] #add item to end of fresh dataframe
                    dfchange.to_csv(logpath+'change.csv', index=False)
                    print(i, item)                
                    print("\nTime taken so far..." + str(datetime.datetime.now() - start))
                    
            
                i = i + 1
            except FileNotFoundError:
                var = input("Network failure: please make sure you're reconnected to the network and press return when ready")



def CreatePage(notes, new_len, del_len, change_len, df_new, df_del, df_final, filters, exportfilepath, menu):
    pd.set_option('display.max_colwidth', -1) #stop the dataframe from truncating cell contents. This needs to be set if you want html links to work in cell contents
    
    html = """
    <head>
    <link rel="stylesheet"
      href="https://stackpath.bootstrapcdn.com/bootstrap/4.3.1/css/bootstrap.min.css"
      integrity="sha384-ggOyR0iXCbMQv3Xipma34MD+dH/1fQ784/j6cY/iJTQUOhcWr7x9JvoRxT2MZw1T"
      crossorigin="anonymous" />
    <link rel="stylesheet" href="https://cdn.datatables.net/1.10.20/css/jquery.dataTables.min.css" />
    
    <link rel="stylesheet" type="text/css" media="screen" />
    <meta http-equiv="content-type" content="text/html; charset=UTF-8" />
    <style>
      td{
          word-wrap: break-word;
          min-width: 160px;
          max-width: 160px;
      }</style>
    <script src="https://code.jquery.com/jquery-3.3.1.js"></script>
    <script src="https://cdn.datatables.net/1.10.20/js/jquery.dataTables.min.js"></script>    
  </head>
    """
    html += '<body><div class="container"><h1 id="home">' + source + ' UPDATES OVERVIEW: <b>' + filters + '</b></h1>' 
    html += menu + '<hr />'
    html += '<p>This is a comparison between the HMRC data sets for July and August 2018.</p>' 
    #html += '<p>Filters: <b>' + filters + '</b></p>' 
    html += '<p>Notes: ' + notes + '</p>'
    html += '<div style="max-width: 800px;">' 
    html += r'<p><b>' + change_len + '</b> changed chapters have been found, <a href="#change">see here</a></p>'
    html += r'<p><b>' + new_len + '</b> new chapters have been added, <a href="#new">see here</a></p>' 
    html += r'<p><b>' + del_len + '</b> chapters have been deleted, <a href="#del">see here</a></p>'
    html += r'</div><div style="max-width: 800px;"><h1 id="change">CHANGES</h1><p>' + change_len + ' chapters have changed and are displayed below</p>'#, (<a href="#home">scroll to top</a>)</p>'
    html += df_final.to_html(justify='left', na_rep = " ",index = False,  classes="table table-bordered text-left table-striped table-hover display")
    html += r'</div><div style="max-width: 800px;"><h1 id="new">ADDITIONS</h1><p>' + new_len + ' new chapters have been added and are displayed below, (<a href="#home">scroll to top</a>)</p>'
    html += df_new.to_html(na_rep = " ",index = False,  classes="table table-bordered text-left table-striped table-hover")
    html += r'</div><div style="max-width: 800px;"><h1 id="del">DELETIONS</h1><p>' + del_len + ' chapters have been deleted and are displayed below, (<a href="#home">scroll to top</a>)</p>'
    html += df_del.to_html(na_rep = " ",index = False,  classes="table table-bordered text-left table-striped table-hover")
    html += """<script>
      $(document).ready(function() {
    $('.dataframe').DataTable();
} );
    </script>"""
    html += '</div></div></body></html>'
    html = html.replace('&lt;', '<').replace('&gt;','>')
    #write to html
    print("Exported HTML to...", exportfilepath)
    with open(exportfilepath,'w', encoding='utf-8') as f:
        f.write(html)
        f.close()
        pass

def Relevant(x):
    if any(dfreference.Pattern == x.ChapterNumber) == True:
        return True
    else:
        return False
def AddRelevance():       
    #Check whether chapter acronym appears on the source list of hmrc
    def RelevantSource(x):
        source = re.search(r'\D+', x.ChapterNumber).group()
        if any(dfreferencesourceonly.Source == source) == True:
            return True
        else:
            return False
        
    dfreference = pd.read_csv(maindir + 'hmrc-patterns-locations-28252019.csv', encoding ='utf-8')

    #Create lookup table of references for each chapter
    dfrelevantChapters = pd.DataFrame()
    relevantChapters = []
    for chapter in dfreference.Pattern:
        relevantChapters.append(chapter)
    dfrelevantChapters['Chapter'] = relevantChapters
    dfrelevantChapters.drop_duplicates(inplace=True)
    DocInfo = []
    i=0
    for index, row in dfrelevantChapters.iterrows():
        Pattern = dfrelevantChapters.Chapter.iloc[i]
        patternCount = len(dfreference[dfreference.Pattern.isin([Pattern])]) #Count number of entries for each chapter number/pattern
        DocList = []
        for x in range(0,patternCount):
            DocTitle = str(dfreference.loc[dfreference['Pattern'] == Pattern, 'DocTitle'].iloc[x])
            DocID = str(dfreference.loc[dfreference['Pattern'] == Pattern, 'DocID'].iloc[x])
            PA = str(dfreference.loc[dfreference['Pattern'] == Pattern, 'PA'].iloc[x])
            Count = str(dfreference.loc[dfreference['Pattern'] == Pattern, 'Count'].iloc[x])
            DocList += [DocID]
        DocInfo += [DocList]
        i=i+1
    
    dfrelevantChapters['DocInfo'] = DocInfo
    dfrelevantChapters = dfrelevantChapters.reset_index()#(drop=True)
    #dfrelevantChapters.to_csv(outputdir + "dfrelevantChapters.csv", sep=',')


    dfreferencesourceonly = pd.DataFrame()
    #Extract source info only from the pattern locations
    search = []    
    for values in dfreference.Pattern:
        search.append(re.search(r'\D+', values).group())

    dfreferencesourceonly['Source'] = search
    dfreferencesourceonly.drop_duplicates(inplace=True)
    dfreferencesourceonly.to_csv(outputdir + "dfreferencesourceonly.csv", sep=',')

def DocInfoGenerator(df, ColName):
    DocInfoList = []
    i=0
    for index, row in df.iterrows():
        ChapterNum = df[ColName].iloc[i]
        #print(ChapterNum)
        DocInfo = ''
        try:
            DocInfo = str(dfrelevantChapters.loc[dfrelevantChapters['Chapter'] == ChapterNum, 'DocInfo'].iloc[0])
        except:
            pass#print('Chapter number not found in list of PSL links to HMRC content...')
        DocInfoList.append(DocInfo)               
        i=i+1
    return DocInfoList


def levenshtein_distance(prev, curr):
    counter = {"+": 0, "-": 0}
    distance = 0
    for edit_code, *_ in ndiff(prev, curr):
        if edit_code == " ":
            distance += max(counter.values())
            counter = {"+": 0, "-": 0}
        else: 
            counter[edit_code] += 1
    distance += max(counter.values())
    return distance


#compare previous and current character lengths of the paras and return difference
def CharacterCountDiff(prev, curr):   
    if prev == curr:
        CharCount = 0
    else:
        if prev > curr: 
            CharCount = prev - curr
        else: 
            if curr > prev: 
                CharCount = curr - prev
            else: CharCount = ''
    #print(CharCount) 
    return CharCount

def LevDistComparison(logpath, xmldir1, xmldir2, NSMAP):
    print("Opening up docs on the change list to calculate Levenshtein Distance...")

    df = pd.DataFrame()
    dfchange = pd.read_csv(logpath+'change.csv')
    i = 0
    for index, item in dfchange.itertuples(): #loop through each item in the remainder list
        #if str(dfchange.iloc[i,0]) == 'cfm26050.xml': 
        filename1 = xmldir1 + str(dfchange.iloc[i,0])
        filename2 = xmldir2 + str(dfchange.iloc[i,0])
        #print(i, filename1, filename2)   

        tree1 = etree.parse(filename1)
        tree2 = etree.parse(filename2)
        page1 = ET.tostring(tree1.getroot(), encoding='utf-8', method='text')
        page2 = ET.tostring(tree2.getroot(), encoding='utf-8', method='text')
        
        chapnumber = tree1.find(".//ci:content", NSMAP).text  
        title = tree1.find(".//docinfo:doc-heading", NSMAP).text
        title = title.replace(chapnumber + ' ', '')
        try: source = tree1.find(".//docinfo:hierlev//docinfo:hierlev//title", NSMAP).text
        except AttributeError: source = 'not found'
        levdisttotal = levenshtein_distance(str(page1), str(page2))
        charcountdiff = CharacterCountDiff(len(str(page1)), len(str(page2)))
        texts1 = tree1.findall(".//text")
        texts2 = tree2.findall(".//text")        

        g=0
        #textlist=[]
        textLevTotal=0
        paraChangeCount=0
        paraCount=0
        if levdisttotal > 0:
            for text1 in texts1:
                #try:                         
                #levdisttemp = levenshtein_distance(str(text1.text), str(texts2[g].text))
                oldtext = ''
                newtext = ''
                try:
                    oldtexts = text1.xpath('.//text()') #list of all text parts within the tag
                    for individualtext in oldtexts: 
                        individualtext = re.sub('\s$', '', individualtext) #remove whitespace from end of string
                        try:oldtext+=etree.strip_tags(individualtext,'remotelink')        
                        except TypeError: oldtext+=individualtext
                    newtexts = texts2[g].xpath('.//text()')
                    for individualtext in newtexts: 
                        individualtext = re.sub('\s$', '', individualtext) #remove whitespace from end of string
                        try: newtext+=etree.strip_tags(individualtext,'remotelink')
                        except TypeError: newtext+=individualtext
                except IndexError: pass
                if oldtext != newtext: paraChangeCount = paraChangeCount + 1
                paraCount+=1
                levdisttemp = levenshtein_distance(oldtext, newtext)
                textLevTotal = textLevTotal + levdisttemp
                #if levdisttemp > levdisttotal:
                #    print(levdisttemp, levdisttotal)
                #    print(oldtext+'\n')
                #    print(newtext)
                #print(str(text1.text)+'\n')
                #print(str(texts2[g].text))
                #if levdisttemp > 0:
                #    levdisttemppercent = round(((levdisttemp / levdisttotal) * 100), 2)
                #    textlist.append(['t'+str(g),levdisttemp, str(levdisttemppercent)+'%'])            
                #except: print('error')
                g+=1
        else:
            textlist = 'NA'
    
        link = 'View changes'
        row = {"Source": source, "ChapterNumber": chapnumber, "Title": title, "LevDistance": levdisttotal, "CharCountDiff": charcountdiff, "textLevTotal": textLevTotal, "paraChangeCount": paraChangeCount, "paraCount": paraCount, "Link":link}
        print(row)
        df = df.append(row, ignore_index=True) 
        df.to_csv(outputdir + "df-text8.csv", sep=',', index=False)
        i+=1


def GatherDocInfo(logpath, xmldir, NSMAP, df, type):
    print("Opening up docs on given list to extract info from each doc...")
    df1 = pd.DataFrame()
    i = 0
    for index, item in df.itertuples(): 
        filename = xmldir + str(df.iloc[i,0])
        if 'xml' in filename:
            tree = etree.parse(filename)
            page = ET.tostring(tree.getroot(), encoding='utf-8', method='text')        
            chapnumber = tree.find(".//ci:content", NSMAP).text  
            title = tree.find(".//docinfo:doc-heading", NSMAP).text
            title = title.replace(chapnumber + ' ', '')
            try: source = tree.find(".//docinfo:hierlev//docinfo:hierlev//title", NSMAP).text
            except AttributeError: source = 'not found'
            link = 'View changes'
            row = {"Source": source, "ChapterNumber": chapnumber, "Title": title, "Link":link}
            print(row)
            df1 = df1.append(row, ignore_index=True) 
            df1.to_csv(outputdir + "df-" + type + ".csv", sep=',', index=False)
        i+=1

    #wait = input("PAUSED...when ready press enter")

def ExportChangeListToHTMLAndDoc(logpath, start, xmldir1, xmldir2, df):    
    i = 0
    print('Stepping through each doc and comparing, outputting to HTML and Word doc...')     
    for item in df.itertuples(): #loop through each item in the remainder list
        #if str(dfchange.iloc[i,0]) == 'cfm26050.xml': 
        #html = '<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.0 Transitional//EN" "http://www.w3.org/TR/xhtml1/DTD/xhtml1-transitional.dtd"><html><head><meta http-equiv="Content-Type" content="text/html; charset=utf-8" /><link rel="stylesheet" type="text/css" href="style.css"/><link rel="stylesheet" href="https://maxcdn.bootstrapcdn.com/bootstrap/4.0.0/css/bootstrap.min.css" integrity="sha384-Gn5384xqQ1aoWXA+058RXPxPg6fy4IWvTNh0E263XmFcJlSAwiGgFAW/dAiS6JXm" crossorigin="anonymous" /><title></title></head><body>'

        filename1 = xmldir1 + str(df['ChapterNumber'].iloc[i] + '.xml')
        filename2 = xmldir2 + str(df['ChapterNumber'].iloc[i] + '.xml')
        #print(i, filename1, filename2)                       
        
        #Import xml file into soup
        tree1 = etree.parse(filename1)
        tree2 = etree.parse(filename2)
        
        etree.strip_tags(tree1,'remotelink')
        etree.strip_tags(tree2,'remotelink')

        #deconstruct xml into dictionaries for comparison later        
        dict1 = {}
        dict2 = {}
        wantedTags = ['text', 'h']
        d=0
        for elem in tree1.iter(): 
            if elem.tag in wantedTags:                
                if elem.text != None:
                    if elem.text != ' ':
                        if elem.tag == 'h': tag = 'h4'
                        else: tag = 'p'                    
                        text = etree.strip_tags(elem,'remotelink')
                        text = re.sub('\s$', '', elem.text) #remove whitespace from end of string
                        dict1.update({d: [tag, text]})
                        d+=1
        
        d=0
        for elem in tree2.iter(): 
            if elem.tag in wantedTags:
                if elem.text != None:
                    if elem.text != ' ':
                        if elem.tag == 'h': tag = 'h4'
                        else: tag = 'p'
                        text = etree.strip_tags(elem,'remotelink')
                        text = re.sub('\s$', '', elem.text) 
                        dict2.update({d: [tag, text]})
                        d+=1

        chapnumber = tree1.find(".//ci:content", NSMAP).text 
        doctitle1 = tree1.find(".//docinfo:doc-heading", NSMAP).text 
        doctitle1 = doctitle1.replace(chapnumber + ' ', '')
        doctitle2 = tree2.find(".//docinfo:doc-heading", NSMAP).text 
        doctitle2 = doctitle2.replace(chapnumber + ' ', '')
        #input1 = tree1.getroot()
        #input2 = tree2.getroot()

        #set up html file to export
        html = etree.Element('html')
        head = etree.SubElement(html, 'head')
        meta = etree.SubElement(head, 'meta')
        meta.set('http-equiv', 'Content-Type')
        meta.set('content', 'text/html; charset=utf-8')
        link = etree.SubElement(head, 'link') #local css
        link.set('rel', 'stylesheet')
        link.set('type', 'text/css')
        link.set('href', '../css/style.css')
        link = etree.SubElement(head, 'link') #bootstrap
        link.set('rel', 'stylesheet')
        link.set('href', 'https://maxcdn.bootstrapcdn.com/bootstrap/4.0.0/css/bootstrap.min.css')
        link.set('itegrity', 'sha384-Gn5384xqQ1aoWXA+058RXPxPg6fy4IWvTNh0E263XmFcJlSAwiGgFAW/dAiS6JXm')
        link.set('crossorigin', 'anonymous')

        title = etree.SubElement(head, 'title')
        title.text = 'HMRC - ' + chapnumber #+ ' - ' + doctitle1
        body = etree.SubElement(html, 'body')
        divmain = etree.SubElement(body, 'div')
        divmain.set('class', 'container')
        h1 = etree.SubElement(divmain, 'h1')
        h1.text = chapnumber #+ ':'
        #h2 = etree.SubElement(divmain, 'h2')
        #h2.text = doctitle1
        hr = etree.SubElement(divmain, 'hr')

        divleft = etree.SubElement(divmain, 'div')
        divleft.set('contenteditable', 'true')
        divleft.set('class', 'container')
        divleft.set('style', 'width: 50%; height: 50%; float:left;')        
        h3 = etree.SubElement(divleft, 'h3')
        h3.text = 'Old version'    
        p = etree.SubElement(divleft, 'p')
        p.text = filename1
        hr = etree.SubElement(divleft, 'hr')
        sup = etree.SubElement(divleft, 'sup')
        sup.text = 'title'
        left = etree.SubElement(divleft, 'h3')
        left.text = doctitle1

        divright = etree.SubElement(divmain, 'div')
        divright.set('contenteditable', 'true')
        divright.set('class', 'container')
        divright.set('style', 'width: 50%; height: 50%; float:right;')   
        h3 = etree.SubElement(divright, 'h3')
        h3.text = 'New version'
        p = etree.SubElement(divright, 'p')
        p.text = filename2
        hr = etree.SubElement(divright, 'hr')
        sup = etree.SubElement(divright, 'sup')
        sup.text = 'title'
        right = etree.SubElement(divright, 'h3')
        right.text = doctitle2
        if doctitle1 != doctitle2: right.set('style', 'background-color: #FFFF00')

        
        #body1 = input1.find(".//pgrp")
        #body2 = input2.find(".//pgrp")
        body2 = tree2.iter()
        x=0
        for entry in dict1:           
            sup = etree.SubElement(divleft, 'sup')
            sup.text = str(entry)
            left = etree.SubElement(divleft, dict1[entry][0])
            left.text = dict1[entry][1]
            try: 
                sup = etree.SubElement(divright, 'sup')
                sup.text = str(entry)
                right = etree.SubElement(divright, dict2[entry][0])
                right.text = dict2[entry][1]
                if dict1[entry][1] != dict2[entry][1]: right.set('style', 'background-color: #FFFF00')
            except KeyError: 
                print('Page 1 has more elements than page 2...')
                x+=1

        #print(ET.tostring(body2))
        

        #html = html.replace("b'<pgrp>",'')
        #html = html.replace("</pgrp>'",'')
        #html = html.replace("<h>",'<h3>')
        #html = html.replace("</h>",'</h3>')
        #html = html.replace("<text>",'')
        #html = html.replace("<text>",'')
        #html = html.replace("</text>",'')
        #html = html.replace("<text />", '')
        #html = html.replace("&#8226;\\t",'')
        #html = html.replace("<row",'<tr')
        #html = html.replace("<entry",'<td')
        #html = html.replace("</row>",'</tr>')
        #html = html.replace("</entry>",'</td>')
        #html = html.replace("<l>",'')
        #html = html.replace("</l>",'')
        #html = html.replace("<p></p>",'')
        #html = html.replace("</lilabel><p>",' ')

        #wait = input("PAUSED...when ready press enter")
        
        
        
        print("Time taken so far..." + str(datetime.datetime.now() - start))
        #if i == 10: break        

        #html += '</body></html>'
        htmlfilepath = outputdir + 'chapters/' + chapnumber + ".html"
        
        tree = etree.ElementTree(html)
        tree.write(htmlfilepath,encoding='utf-8')

        #with open(htmlfilepath,'w', encoding='utf-8') as f:
        #    f.write(html)  #select the first table and append to the html variable
        #    f.close()
        #    pass

        print("Exported to: " + htmlfilepath)

        
        def XMLtoDoc(dictionary, tree, doctitle, chapnumber, outputdir, type):
            document = Document()
            document.add_heading('HMRC - ' + chapnumber + ' - ' + doctitle, 0)
            for entry in dictionary: document.add_paragraph(dictionary[entry][1])            
            docfilepath = outputdir + chapnumber + '-' + type + ".docx"    
            document.save(docfilepath) 
            #print("Exported to: " + docfilepath)

        
        
        XMLtoDoc(dict1, tree1, doctitle1, chapnumber, outputdir + '\\Chapters\\', 'old')    
        XMLtoDoc(dict2, tree2, doctitle2, chapnumber, outputdir + '\\Chapters\\', 'new')    
        #Create the Application word
        Application=win32com.client.gencache.EnsureDispatch("Word.Application")
        Application.CompareDocuments(Application.Documents.Open('C:\\GIT\\HMRC\\www\\chapters\\' + chapnumber + '-old.docx'), Application.Documents.Open('C:\\GIT\\HMRC\\www\\chapters\\' + chapnumber + '-new.docx'))
        #Application.ActiveDocument.ActiveWindow.View.Type = 3 # prevent that word opens itself
        Application.ActiveDocument.SaveAs (FileName = 'C:\\GIT\\HMRC\\www\\chapters\\' + chapnumber + '-comp.docx')
        Application.Documents.Close()
        Application.Quit()
        #deleteList.append('C:\\GIT\\HMRC\\www\\chapters\\' + chapnumber + "-old.docx")
        os.remove('C:\\GIT\\HMRC\\www\\chapters\\' + chapnumber + "-old.docx")
        #deleteList.append('C:\\GIT\\HMRC\\www\\chapters\\' + chapnumber + "-new.docx")
        os.remove('C:\\GIT\\HMRC\\www\\chapters\\' + chapnumber + "-new.docx")

        print("Exported to: " + 'C:\\GIT\\HMRC\\www\\chapters\\' + chapnumber + '-comp.docx')


        i = i + 1
    

def ExportListToHTMLAndDoc(logpath, start, xmldir, df):    
    i = 0
    print('Stepping through each doc and building simple HTML version...')     
    for item in df.itertuples(): 
        filename = xmldir + str(df['ChapterNumber'].iloc[i] + '.xml')
        tree = etree.parse(filename)
        etree.strip_tags(tree,'remotelink')
        #deconstruct xml into dictionaries for comparison later        
        dictionary = {}
        wantedTags = ['text', 'h']
        d=0
        for elem in tree.iter(): 
            if elem.tag in wantedTags:                
                if elem.text != None:
                    if elem.text != ' ':
                        if elem.tag == 'h': tag = 'h4'
                        else: tag = 'p'                    
                        text = etree.strip_tags(elem,'remotelink')
                        text = re.sub('\s$', '', elem.text) #remove whitespace from end of string
                        dictionary.update({d: [tag, text]})
                        d+=1        
        
        chapnumber = tree.find(".//ci:content", NSMAP).text 
        doctitle = tree.find(".//docinfo:doc-heading", NSMAP).text 
        doctitle = doctitle.replace(chapnumber + ' ', '')

        def UnpackToHTML(dictionary, tree, doctitle, chapnumber, outputdir):
            html = etree.Element('html')
            head = etree.SubElement(html, 'head')
            meta = etree.SubElement(head, 'meta')
            meta.set('http-equiv', 'Content-Type')
            meta.set('content', 'text/html; charset=utf-8')
            link = etree.SubElement(head, 'link') #local css
            link.set('rel', 'stylesheet')
            link.set('type', 'text/css')
            link.set('href', '../css/style.css')
            link = etree.SubElement(head, 'link') #bootstrap
            link.set('rel', 'stylesheet')
            link.set('href', 'https://maxcdn.bootstrapcdn.com/bootstrap/4.0.0/css/bootstrap.min.css')
            link.set('itegrity', 'sha384-Gn5384xqQ1aoWXA+058RXPxPg6fy4IWvTNh0E263XmFcJlSAwiGgFAW/dAiS6JXm')
            link.set('crossorigin', 'anonymous')

            title = etree.SubElement(head, 'title')
            title.text = 'HMRC - ' + chapnumber + ' - ' + doctitle
            body = etree.SubElement(html, 'body')
            divmain = etree.SubElement(body, 'div')
            divmain.set('class', 'container')
            h1 = etree.SubElement(divmain, 'h1')
            h1.text = chapnumber #+ ':'
            hr = etree.SubElement(divmain, 'hr')
            title = etree.SubElement(divmain, 'h3')
            title.text = doctitle
            
            for entry in dictionary:           
                sup = etree.SubElement(divmain, 'sup')
                sup.text = str(entry)
                line = etree.SubElement(divmain, dictionary[entry][0])
                line.text = dictionary[entry][1]
                        
            htmlfilepath = outputdir + chapnumber + ".html"        
            tree = etree.ElementTree(html)
            tree.write(htmlfilepath,encoding='utf-8')
            print("Exported to HTML page: " + htmlfilepath)        
        
        UnpackToHTML(dictionary, tree, doctitle, chapnumber, outputdir + '\\Chapters\\')   

        def UnpackToDoc(dictionary, tree, doctitle, chapnumber, outputdir):
            document = Document()
            document.add_heading('HMRC - ' + chapnumber + ' - ' + doctitle, 0)
            for entry in dictionary: document.add_paragraph(dictionary[entry][1])            
            docfilepath = outputdir + chapnumber + ".docx"    
            document.save(docfilepath) 
            print("Exported to Word doc: " + docfilepath)

        UnpackToDoc(dictionary, tree, doctitle, chapnumber, outputdir + '\\Chapters\\') 
        
        i = i + 1




def ExportChangeListToDoc(logpath, start, xmldir1, xmldir2, df):    
    i = 0
    print('Stepping through each doc and comparing...')     
    for item in df.itertuples(): #loop through each item in the remainder list
        #if str(dfchange.iloc[i,0]) == 'cfm26050.xml': 
        #html = '<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.0 Transitional//EN" "http://www.w3.org/TR/xhtml1/DTD/xhtml1-transitional.dtd"><html><head><meta http-equiv="Content-Type" content="text/html; charset=utf-8" /><link rel="stylesheet" type="text/css" href="style.css"/><link rel="stylesheet" href="https://maxcdn.bootstrapcdn.com/bootstrap/4.0.0/css/bootstrap.min.css" integrity="sha384-Gn5384xqQ1aoWXA+058RXPxPg6fy4IWvTNh0E263XmFcJlSAwiGgFAW/dAiS6JXm" crossorigin="anonymous" /><title></title></head><body>'

        filename1 = xmldir1 + str(df['ChapterNumber'].iloc[i] + '.xml')
        filename2 = xmldir2 + str(df['ChapterNumber'].iloc[i] + '.xml')
        #print(i, filename1, filename2)                       
        
        #Import xml file into soup
        tree1 = etree.parse(filename1)
        tree2 = etree.parse(filename2)
        
        etree.strip_tags(tree1,'remotelink')
        etree.strip_tags(tree2,'remotelink')

        #deconstruct xml into dictionaries for comparison later        
        dict1 = {}
        dict2 = {}
        wantedTags = ['text', 'h']
        d=0
        for elem in tree1.iter(): 
            if elem.tag in wantedTags:                
                if elem.text != None:
                    if elem.text != ' ':
                        if elem.tag == 'h': tag = 'h4'
                        else: tag = 'p'                    
                        text = etree.strip_tags(elem,'remotelink')
                        text = re.sub('\s$', '', elem.text) #remove whitespace from end of string
                        dict1.update({d: [tag, text]})
                        d+=1
        
        d=0
        for elem in tree2.iter(): 
            if elem.tag in wantedTags:
                if elem.text != None:
                    if elem.text != ' ':
                        if elem.tag == 'h': tag = 'h4'
                        else: tag = 'p'
                        text = etree.strip_tags(elem,'remotelink')
                        text = re.sub('\s$', '', elem.text) 
                        dict2.update({d: [tag, text]})
                        d+=1
        
        chapnumber = tree1.find(".//ci:content", NSMAP).text 
        doctitle1 = tree1.find(".//docinfo:doc-heading", NSMAP).text 
        doctitle1 = doctitle1.replace(chapnumber + ' ', '')
        doctitle2 = tree2.find(".//docinfo:doc-heading", NSMAP).text 
        doctitle2 = doctitle2.replace(chapnumber + ' ', '')

        def XMLtoDoc(dictionary, tree, doctitle, chapnumber, outputdir, type):
            document = Document()
            document.add_heading('HMRC - ' + chapnumber + ' - ' + doctitle, 0)
            for entry in dictionary: document.add_paragraph(dictionary[entry][1])            
            docfilepath = outputdir + chapnumber + '-' + type + ".docx"    
            document.save(docfilepath) 
            #print("Exported to: " + docfilepath)

        
        
        XMLtoDoc(dict1, tree1, doctitle1, chapnumber, outputdir + '\\Chapters\\', 'old')    
        XMLtoDoc(dict2, tree2, doctitle2, chapnumber, outputdir + '\\Chapters\\', 'new')    
        #Create the Application word
        Application=win32com.client.gencache.EnsureDispatch("Word.Application")
        Application.CompareDocuments(Application.Documents.Open('C:\\GIT\\HMRC\\www\\chapters\\' + chapnumber + '-old.docx'), Application.Documents.Open('C:\\GIT\\HMRC\\www\\chapters\\' + chapnumber + '-new.docx'))
        #Application.ActiveDocument.ActiveWindow.View.Type = 3 # prevent that word opens itself
        Application.ActiveDocument.SaveAs (FileName = 'C:\\GIT\\HMRC\\www\\chapters\\' + chapnumber + '-comp.docx')
        Application.Documents.Close()
        Application.Quit()
        deleteList.append('C:\\GIT\\HMRC\\www\\chapters\\' + chapnumber + "-old.docx")
        deleteList.append('C:\\GIT\\HMRC\\www\\chapters\\' + chapnumber + "-new.docx")
        print("Exported to: " + 'C:\\GIT\\HMRC\\www\\chapters\\' + chapnumber + '-comp.docx')
        i = i + 1
    




#### MAIN ####

date =  time.strftime("%d/%m/%Y")
logdate = datetime.datetime.now()
logname = time.strftime("HMRC-02FO-%d%m%Y-%H%M%S.csv")
#xmldir1 =  "C:\\Users\\Hutchida\\Documents\\HMRC\\july\\"
#xmldir2 =  "C:\\Users\\Hutchida\\Documents\\HMRC\\august\\"
xmldir1 = '\\\\lngoxfdatp16vb\\Fabrication\\Build\\0Y02\Data\\'
xmldir2 = '\\\\lngoxfdatp16vb\\Fabrication\\Build\\0Y04\\Data\\'
logpath = "C:\\Users\\Hutchida\\Documents\\HMRC\\logs\\"
source = "HMRC"
maindir = "C:\\Users\\Hutchida\\Documents\\HMRC\\"
#outputdir = maindir + "\\output\\"
outputdir = "www\\"
NSMAP = {'lnvxe': 'http://www.lexis-nexis.com/lnvxe', 'lnv': 'http://www.lexis-nexis.com/lnv', 'lnvni': 'http://www.lexis-nexis.com/lnvni', 'lnclx': 'http://www.lexis-nexis.com/lnclx', 'lncle': 'http://www.lexis-nexis.com/lncle', 'lndel': 'http://www.lexis-nexis.com/lndel', 'lngntxt': 'http://www.lexis-nexis.com/lngntxt', 'lndocmeta': 'http://www.lexis-nexis.com/lndocmeta', 'lnlit': 'http://www.lexis-nexis.com/lnlit', 'lnci': 'http://www.lexis-nexis.com/lnci', 'nitf': 'urn:nitf:iptc.org.20010418.NITF', 'lnvx': 'http://www.lexis-nexis.com/lnvx', 'ci': 'http://www.lexis-nexis.com/ci', 'glp': 'http://www.lexis-nexis.com/glp', 'case': 'http://www.lexis-nexis.com/glp/case', 'jrnl': 'http://www.lexis-nexis.com/glp/jrnl', 'comm': 'http://www.lexis-nexis.com/glp/comm', 'cttr': 'http://www.lexis-nexis.com/glp/cttr', 'dict': 'http://www.lexis-nexis.com/glp/dict', 'dig': 'http://www.lexis-nexis.com/glp/dig', 'docinfo': 'http://www.lexis-nexis.com/glp/docinfo', 'frm': 'http://www.lexis-nexis.com/glp/frm', 'in': 'http://www.lexis-nexis.com/glp/in', 'leg': 'http://www.lexis-nexis.com/glp/leg', 'xhtml': 'http://www.w3c.org/1999/xhtml'}
deleteList = []
start = datetime.datetime.now() 

#print("\nComparing directories to discover additions and deletions..." + str(datetime.datetime.now()))
#DirectoryComparison(logpath, xmldir1, xmldir2)

#print("\nEntering prelim phase..." + str(datetime.datetime.now()))
#Prelim(xmldir1, xmldir2, logpath, start)

#print("\n\n\nSo far time taken..." + str(datetime.datetime.now() - start))

#reset the log date the filename will take so that it differs from the previous log
#logdate = datetime.datetime.now()
#print("\n\n\nEntering Lev Dist comparison phase..." + str(datetime.datetime.now()))

#Build detailed lists
#LevDistComparison(logpath, xmldir1, xmldir2, NSMAP)
#GatherDocInfo(logpath, xmldir2, NSMAP, pd.read_csv(outputdir + "additions.csv"), 'additions')
#GatherDocInfo(logpath, xmldir1, NSMAP, pd.read_csv(outputdir + "deletions.csv"), 'deletions')
#wait = input("PAUSED...when ready press enter")


df_new = pd.read_csv(outputdir + "df-additions.csv")
df_del = pd.read_csv(outputdir + "df-deletions.csv")
df_final = pd.read_csv(outputdir + "df-text8.csv")

#tidy up df_final a bit, stick this in a function later on
#remove link column and rebuild it
del df_final['Link']
df_final['HTML'] = '<a href="chapters/' + df_final['ChapterNumber'] + '.html" target="_blank">View</a>' #mark up the link ready for html
df_final.HTML = df_final.HTML.str.replace('—', ' ')
df_final.HTML = df_final.HTML.str.replace('’',"'")
df_final = df_final[(df_final['paraChangeCount'] != 0)]
df_final = df_final[(df_final['LevDistance'] > 50)]
df_final['TrackedChanges'] = '<a href="chapters/' + df_final['ChapterNumber'] + '-comp.docx" target="_blank">Download</a>' 

df_final.LevDistance = df_final.LevDistance.astype(int)
df_final.textLevTotal = df_final.textLevTotal.astype(int)
df_final.CharCountDiff = df_final.CharCountDiff.astype(int)
df_final.paraCount = df_final.paraCount.astype(int)
df_final.paraChangeCount = df_final.paraChangeCount.astype(int)

df_final['ChangeInfo'] = 'Lev Distance: ' +  df_final.LevDistance.astype(str) + '<br />CharCountDiff: ' + df_final.CharCountDiff.astype(str) + '<br />Changed paras: ' + df_final.paraChangeCount.astype(str) + ' out of ' + df_final.paraCount.astype(str) 
df_final = df_final.sort_values(['LevDistance'], ascending = False)
columnsTitles = ['Source', 'ChapterNumber', 'Title', 'ChangeInfo', 'HTML', 'TrackedChanges']
df_final = df_final.reindex(columns=columnsTitles)

df_new['HTML'] = '<a href="chapters/' + df_new['ChapterNumber'] + '.html" target="_blank">View</a>' #mark up the link ready for html
df_new.HTML = df_new.HTML.str.replace('—', ' ')
df_new.HTML = df_new.HTML.str.replace('’',"'")
df_new['TrackedChanges'] = '<a href="chapters/' + df_new['ChapterNumber'] + '.docx" target="_blank">Download</a>' 
df_new = df_new.sort_values(['ChapterNumber'], ascending = False)
columnsTitles = ['Source', 'ChapterNumber', 'Title', 'HTML', 'TrackedChanges']
df_new = df_new.reindex(columns=columnsTitles)

df_del['HTML'] = '<a href="chapters/' + df_del['ChapterNumber'] + '.html" target="_blank">View</a>' #mark up the link ready for html
df_del.HTML = df_del.HTML.str.replace('—', ' ')
df_del.HTML = df_del.HTML.str.replace('’',"'")
df_del['TrackedChanges'] = '<a href="chapters/' + df_del['ChapterNumber'] + '.docx" target="_blank">Download</a>' 
df_del = df_del.sort_values(['ChapterNumber'], ascending = False)
columnsTitles = ['Source', 'ChapterNumber', 'Title', 'HTML', 'TrackedChanges']
df_del = df_del.reindex(columns=columnsTitles)


#Generate HTML and Word chapter docs
print('ExportListToHTMLAndDoc...for additions list')
#ExportListToHTMLAndDoc(logpath, start, xmldir2, df_new) 
print('ExportListToHTMLAndDoc...for deletions list')  
#ExportListToHTMLAndDoc(logpath, start, xmldir1, df_del)   
print('ExportChangeListToHTMLAndDoc...for change list')  
ExportChangeListToHTMLAndDoc(logpath, start, xmldir1, xmldir2, df_final)



menu = '<p><a href="index.html">All</a> | <a href="tax.html">Tax</a> | <a href="privateclient.html">Private Client</a> | <a href="shareschemes.html">Share Schemes</a></p>'
#menu = ''
notes = 'Changed pages only showing below. Added and deleted pages will come shortly. Sorted by Levenshtein Distance to bring the docs with the most changes to the top of the list. In order to give the Lev Distance context, have also added a column that counts how many paras within the doc are not equal, and another column to show how total number of paras in that document.'
CreatePage(notes, str(len(df_new)), str(len(df_del)), str(len(df_final)), df_new, df_del, df_final, 'ALL', outputdir + 'index.html', menu)

notes='Cross referenced with a list of sources that are relevant to the Tax team.'
dfList = pd.read_csv('tax.csv')
taxList = dfList['Source'].values.tolist()
dfTax = df_final[df_final['Source'].isin(taxList)]
dfTaxNew = df_new[df_new['Source'].isin(taxList)]
dfTaxDel = df_del[df_del['Source'].isin(taxList)]

CreatePage(notes, str(len(dfTaxNew)), str(len(dfTaxDel)), str(len(dfTax)), dfTaxNew, dfTaxDel, dfTax, 'TAX RELEVANT', outputdir + 'tax.html', menu)



notes='Cross referenced with a list of sources that are relevant to the Private Client team.'
dfList = pd.read_csv('privateclient.csv')
PCList = dfList['Source'].values.tolist()
dfPC = df_final[df_final['Source'].isin(PCList)]
dfPCNew = df_new[df_new['Source'].isin(PCList)]
dfPCDel = df_del[df_del['Source'].isin(PCList)]

CreatePage(notes, str(len(dfPCNew)), str(len(dfPCDel)), str(len(dfPC)), dfPCNew, dfPCDel, dfPC, 'PRIVATE CLIENT RELEVANT', outputdir + 'privateclient.html', menu)

notes='Cross referenced with a list of sources that are relevant to the Share Schemes team.'
dfList = pd.read_csv('shareschemes.csv')
SSList = dfList['Source'].values.tolist()
dfSS = df_final[df_final['Source'].isin(SSList)]
dfSSNew = df_new[df_new['Source'].isin(SSList)]
dfSSDel = df_del[df_del['Source'].isin(SSList)]

CreatePage(notes, str(len(dfSSNew)), str(len(dfSSDel)), str(len(dfSS)), dfSSNew, dfSSDel, dfSS, 'SHARE SCHEMES RELEVANT', outputdir + 'shareschemes.html', menu)


#ExportToDoc(logpath, start, xmldir1, xmldir2, df_final)   

print("Time taken so far..." + str(datetime.datetime.now() - start))

#for doc in deleteList: os.remove(doc)

print("\n\n\nFinished! Total time taken..." + str(datetime.datetime.now() - start))

#Export(logpath)
#wait = input("PAUSED...when ready press enter")